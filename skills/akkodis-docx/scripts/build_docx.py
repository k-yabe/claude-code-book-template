#!/usr/bin/env python3
"""AKKODiS ブランド準拠 DOCX を生成する。

入力 (Markdown ライク):
    # 種別: proposal | minutes | memo | report
    # タイトル: ...
    # 宛先: ...
    # 作成: ...
    # 機密: Internal | Confidential   （任意・省略時 Internal）

    ## 見出し1
    本文 / 箇条書き / Markdown 表

    ### 小見出し
    本文

依存: python-docx
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from notation import correct  # noqa: E402

NAVY = RGBColor(0x00, 0x1F, 0x33)
GOLD = RGBColor(0xFF, 0xB8, 0x1C)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x5A, 0x64, 0x70)


@dataclass
class Block:
    kind: str  # "para" | "bullet" | "number" | "table"
    text: str = ""
    items: list[str] = field(default_factory=list)
    table: list[list[str]] = field(default_factory=list)


@dataclass
class Section:
    title: str
    level: int  # 1 or 2 (heading depth)
    blocks: list[Block] = field(default_factory=list)


@dataclass
class Doc:
    kind: str = "proposal"
    title: str = "Untitled"
    audience: str = ""
    author: str = ""
    confidential: str = "Internal"
    sections: list[Section] = field(default_factory=list)


def parse_input(text: str) -> Doc:
    doc = Doc()
    current: Section | None = None
    current_block: Block | None = None
    in_table = False
    table_buffer: list[str] = []

    def flush_block() -> None:
        nonlocal current_block
        if current_block is not None and current is not None:
            current.blocks.append(current_block)
            current_block = None

    def flush_table() -> None:
        nonlocal in_table, table_buffer
        if not table_buffer or current is None:
            table_buffer = []
            in_table = False
            return
        rows = [r for r in table_buffer if r.strip()]
        table_buffer = []
        in_table = False
        if len(rows) < 2:
            return

        def split_cells(line: str) -> list[str]:
            line = line.strip()
            if line.startswith("|"):
                line = line[1:]
            if line.endswith("|"):
                line = line[:-1]
            return [correct(c.strip()) for c in line.split("|")]

        header = split_cells(rows[0])
        body_rows = rows[2:] if re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", rows[1]) else rows[1:]
        table = [header] + [split_cells(r) for r in body_rows]
        flush_block()
        current.blocks.append(Block(kind="table", table=table))

    meta_re = {
        r"^#\s*(?:種別|kind)\s*[:：]\s*(.+)$": "kind",
        r"^#\s*(?:タイトル|title)\s*[:：]\s*(.+)$": "title",
        r"^#\s*(?:宛先|audience)\s*[:：]\s*(.+)$": "audience",
        r"^#\s*(?:作成|author)\s*[:：]\s*(.+)$": "author",
        r"^#\s*(?:機密|confidential)\s*[:：]\s*(.+)$": "confidential",
    }

    for raw in text.splitlines():
        line = raw.rstrip()
        if line.lstrip().startswith("|"):
            flush_block()
            in_table = True
            table_buffer.append(line)
            continue
        if in_table and not line.strip():
            flush_table()
            continue
        if in_table:
            flush_table()

        # メタ
        if current is None:
            matched = False
            for pattern, attr in meta_re.items():
                m = re.match(pattern, line, re.IGNORECASE)
                if m:
                    val = m.group(1).strip()
                    setattr(doc, attr, val.lower() if attr == "kind" else val)
                    matched = True
                    break
            if matched:
                continue

        # 見出し
        m = re.match(r"^(##+)\s+(.+)$", line)
        if m:
            flush_block()
            level = min(2, len(m.group(1)) - 1)  # ##=1, ###+=2
            current = Section(title=correct(m.group(2).strip()), level=level)
            doc.sections.append(current)
            continue

        if current is None:
            continue

        # 箇条書き
        m = re.match(r"^\s*[-*]\s+(.+)$", line)
        if m:
            if current_block is None or current_block.kind != "bullet":
                flush_block()
                current_block = Block(kind="bullet")
            current_block.items.append(correct(m.group(1).strip()))
            continue
        m = re.match(r"^\s*([0-9]+)\.\s+(.+)$", line)
        if m:
            if current_block is None or current_block.kind != "number":
                flush_block()
                current_block = Block(kind="number")
            current_block.items.append(correct(m.group(2).strip()))
            continue

        if not line.strip():
            flush_block()
            continue

        # 段落
        if current_block is None or current_block.kind != "para":
            flush_block()
            current_block = Block(kind="para")
        current_block.text = (current_block.text + ("\n" if current_block.text else "") + correct(line.strip()))

    flush_table()
    if current_block is not None and current is not None:
        current.blocks.append(current_block)
    return doc


def _set_run_font(run, *, size_pt, bold=False, color: RGBColor | None = None, font="Noto Sans JP"):
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def _shade_paragraph(paragraph, hex_rgb: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_rgb)
    pPr.append(shd)


def _shade_cell(cell, hex_rgb: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_rgb)
    tcPr.append(shd)


def _build_styles(document: Document) -> None:
    """AKKODiS Heading 1 / Heading 2 / Body スタイルを定義する。"""
    styles = document.styles

    def ensure_style(name: str):
        if name in styles:
            return styles[name]
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    h1 = styles["Heading 1"]
    h1.font.name = "Noto Sans JP"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Noto Sans JP"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    normal = styles["Normal"]
    normal.font.name = "Noto Sans JP"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)


def _add_field(paragraph, instr_text: str) -> None:
    """段落に Word フィールド（{ TOC \\o "1-2" \\h \\z \\u } など）を挿入する。"""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def add_cover(document: Document, doc: Doc) -> None:
    # 上部の Navy 帯（パラグラフ shading で実現）
    for _ in range(3):
        p = document.add_paragraph()
        _shade_paragraph(p, "001F33")
        run = p.add_run(" ")
        _set_run_font(run, size_pt=18, color=NAVY)

    # タイトル
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(correct(doc.title))
    _set_run_font(run, size_pt=32, bold=True, color=NAVY)

    # サブメタ（宛先・作成・日付）
    if doc.audience:
        p = document.add_paragraph()
        run = p.add_run(correct(doc.audience))
        _set_run_font(run, size_pt=14, color=BLACK)

    if doc.author:
        p = document.add_paragraph()
        run = p.add_run("作成: " + correct(doc.author))
        _set_run_font(run, size_pt=11, color=GRAY)

    p = document.add_paragraph()
    run = p.add_run(f"日付: {date.today().strftime('%Y年%m月%d日')}")
    _set_run_font(run, size_pt=11, color=GRAY)

    # 機密区分
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"AKKODiS / {correct(doc.confidential)}")
    _set_run_font(run, size_pt=10, color=GRAY)

    # 改ページ
    p = document.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    run = p.add_run("目次")
    _set_run_font(run, size_pt=18, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(8)

    p = document.add_paragraph()
    _add_field(p, 'TOC \\o "1-2" \\h \\z \\u')

    # ToC 後に改ページ
    p = document.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_block(document: Document, block: Block) -> None:
    if block.kind == "para":
        p = document.add_paragraph(style="Normal")
        run = p.add_run(block.text)
        _set_run_font(run, size_pt=11, color=BLACK)
    elif block.kind == "bullet":
        for item in block.items:
            try:
                p = document.add_paragraph(style="List Bullet")
            except KeyError:
                p = document.add_paragraph()
            run = p.add_run(correct(item))
            _set_run_font(run, size_pt=11, color=BLACK)
    elif block.kind == "number":
        for item in block.items:
            try:
                p = document.add_paragraph(style="List Number")
            except KeyError:
                p = document.add_paragraph()
            run = p.add_run(correct(item))
            _set_run_font(run, size_pt=11, color=BLACK)
    elif block.kind == "table":
        if not block.table:
            return
        rows = block.table
        ncols = max(len(r) for r in rows)
        table = document.add_table(rows=len(rows), cols=ncols)
        table.style = "Light Grid Accent 1"
        for ri, row in enumerate(rows):
            for ci in range(ncols):
                cell = table.cell(ri, ci)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                text = row[ci] if ci < len(row) else ""
                # 既存パラグラフを使う
                cell.paragraphs[0].text = ""
                run = cell.paragraphs[0].add_run(text)
                if ri == 0:
                    _set_run_font(run, size_pt=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
                    _shade_cell(cell, "001F33")
                else:
                    _set_run_font(run, size_pt=11, color=BLACK)


def add_section(document: Document, section: Section) -> None:
    # 見出しレベル1は改ページ前置（表紙・ToC 直後を除き）
    style_name = "Heading 1" if section.level == 1 else "Heading 2"
    h = document.add_paragraph(style=style_name)
    run = h.add_run(correct(section.title))
    _set_run_font(
        run,
        size_pt=18 if section.level == 1 else 14,
        bold=True,
        color=NAVY,
    )

    for block in section.blocks:
        add_block(document, block)


def configure_page(document: Document, doc: Doc) -> None:
    for sec in document.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
        # ヘッダー
        header = sec.header
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()
        hp.text = ""
        run = hp.add_run(correct(doc.title))
        _set_run_font(run, size_pt=9, color=NAVY)
        # フッター: 中央 Confidential、右ページ番号
        footer = sec.footer
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(f"AKKODiS / {correct(doc.confidential)}")
        _set_run_font(run, size_pt=9, color=GRAY)
        # ページ番号 (右に追加)
        run_pg = fp.add_run("\t\t")
        _set_run_font(run_pg, size_pt=9, color=GRAY)
        _add_field(fp, "PAGE")


def build(doc: Doc) -> Document:
    document = Document()
    _build_styles(document)
    configure_page(document, doc)
    add_cover(document, doc)
    if doc.kind == "proposal" and len(doc.sections) >= 3:
        add_toc(document)
    for s in doc.sections:
        add_section(document, s)
    return document


def main(argv: list[str]) -> int:
    ap = argparse.ArgumentParser(description="AKKODiS ブランド準拠 DOCX 生成")
    ap.add_argument("--input", type=Path, required=True)
    ap.add_argument("--output", type=Path, required=True)
    args = ap.parse_args(argv)
    doc = parse_input(args.input.read_text(encoding="utf-8"))
    document = build(doc)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(args.output))
    print(f"[ok] {args.output}  (sections: {len(doc.sections)})")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
